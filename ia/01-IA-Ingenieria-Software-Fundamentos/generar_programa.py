from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = 'Arial'
FONT_SIZE = Pt(12)
FONT_SIZE_PT = 152400
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x1b, 0x1c, 0x1d)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_empty_line(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = True
    run.font.color.rgb = BLACK
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = True
    run.font.color.rgb = BLACK
    return p


def add_content(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = BLACK
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = False
    run.font.color.rgb = BLACK
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = False
    run.font.color.rgb = BLACK
    return p


def add_annex_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = True
    run.font.color.rgb = BLACK
    return p


def create_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.font.bold = True
                run.font.color.rgb = BLACK
        set_cell_shading(cell, 'D9E2F3')
    
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE
                    run.font.bold = False
                    run.font.color.rgb = BLACK
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width
    
    return table


def create_programa():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    add_empty_line(doc)
    add_title(doc, 'Programa de')
    add_title(doc, 'IA para la Ingeniería de Software')
    add_empty_line(doc, 3)
    
    add_section_title(doc, '1. NOMBRE DE LA UNIDAD CURRICULAR')
    add_empty_line(doc)
    add_content(doc, 'IA para la Ingeniería de Software')
    add_empty_line(doc, 2)
    
    add_section_title(doc, '2. CRÉDITOS')
    add_empty_line(doc)
    add_content(doc, '8 créditos')
    add_empty_line(doc, 2)
    
    add_section_title(doc, '3. OBJETIVOS DE LA UNIDAD CURRICULAR')
    add_empty_line(doc)
    add_bullet(doc, 'Adquirir los conceptos fundamentales de la IA, desde redes neuronales básicas hasta arquitecturas Transformer')
    add_bullet(doc, 'Implementar el patrón Memory Bank para la gestión de contexto en proyectos de software')
    add_bullet(doc, 'Dominar herramientas de desarrollo asistido (Cursor, Roo Code) y orquestadores básicos (LangChain)')
    add_bullet(doc, 'Desarrollar habilidades prácticas en programación de agentes con patrones ReAct')
    add_bullet(doc, 'Aplicar principios de ingeniería de contexto y automatización de requisitos')
    add_empty_line(doc, 2)
    
    add_section_title(doc, '4. METODOLOGÍA DE ENSEÑANZA')
    add_empty_line(doc)
    
    add_content(doc, 'Participación de los Estudiantes:', bold_prefix='')
    add_empty_line(doc)
    add_bullet(doc, 'Demostraciones Prácticas: El instructor realizará demostraciones en el laboratorio de herramientas de desarrollo asistido por IA durante las horas de clase')
    add_bullet(doc, 'Presentaciones Grupales (Parcial 1): Los estudiantes realizarán presentaciones de 25 minutos en grupos sobre fundamentos de IA')
    add_bullet(doc, 'Proyecto Integrador (Evaluación Final): Desarrollo de un proyecto de aplicación de IA en ingeniería de software')
    add_empty_line(doc, 2)
    
    add_section_title(doc, '5. TEMARIO')
    add_empty_line(doc)
    
    headers = ['Tema', 'Descripción y subtemas']
    rows = [
        ['TEMA I: Fundamentos de IA', 'Historia de la IA, Perceptrón, Funciones de Activación, Algoritmos de Búsqueda A*'],
        ['TEMA II: Deep Learning', 'Redes Neuronales Profundas, Mecanismo de Atención, Transformers'],
        ['TEMA III: Agentes y Orquestación', 'Programación de Agentes, Patrones ReAct, LangChain'],
        ['TEMA IV: Ingeniería de Contexto', 'Memory Bank, Automatización, Pruebas Asistidas'],
    ]
    create_table_with_header(doc, headers, rows, [Cm(4), Cm(13)])
    add_empty_line(doc)
    
    headers_weekly = ['Semana', 'Temas (4 hs de clase)', 'Foco Práctico (2 hs)']
    rows_weekly = [
        ['Semana 1', 'TEMA I: Introducción a la IA', 'Demo: Configuración de entorno de desarrollo con IA'],
        ['Semana 2', 'TEMA I: Perceptrón y Funciones de Activación', 'Demo: Implementación de perceptrón básico'],
        ['Semana 3', 'TEMA I: Algoritmos de Búsqueda A*', 'Demo: Resolución de problemas con A*'],
        ['Semana 4', 'Evaluación Parcial 1 - Presentaciones Grupales', 'Preparación de presentaciones'],
        ['Semana 5', 'TEMA II: Redes Neuronales Profundas', 'Demo: Implementación de red neuronal con PyTorch'],
        ['Semana 6', 'TEMA II: Mecanismo de Atención', 'Demo: Visualización de atención'],
        ['Semana 7', 'TEMA II: Arquitectura Transformer', 'Demo: Uso de modelos pre-entrenados'],
        ['Semana 8', 'TEMA II: Transformers en detalle', 'Demo: Fine-tuning de modelos'],
        ['Semana 9', 'TEMA III: Programación de Agentes', 'Demo: Creación de agentes básicos'],
        ['Semana 10', 'TEMA III: Patrones ReAct', 'Demo: Implementación de patrón ReAct'],
        ['Semana 11', 'TEMA III: LangChain y orquestadores', 'Demo: Cadena de herramientas con LangChain'],
        ['Semana 12', 'TEMA III: Integración y despliegue', 'Demo: Despliegue de agentes'],
        ['Semana 13', 'TEMA IV: Memory Bank y gestión de contexto', 'Demo: Implementación de Memory Bank'],
        ['Semana 14', 'TEMA IV: Automatización de requisitos', 'Demo: Generación automatizada de specs'],
        ['Semana 15', 'TEMA IV: Pruebas asistidas por IA', 'Demo: Generación automática de tests'],
        ['Semana 16', 'Evaluación Final - Proyecto Integrador', 'Defensa de proyectos'],
    ]
    create_table_with_header(doc, headers_weekly, rows_weekly, [Cm(2.5), Cm(6), Cm(8.5)])
    add_empty_line(doc)
    
    add_section_title(doc, '6. BIBLIOGRAFÍA')
    add_empty_line(doc)
    
    add_section_title(doc, '6.1 Básica')
    add_empty_line(doc)
    add_content(doc, 'Goodfellow, I., Bengio, Y., Courville, A. (2016). Deep Learning. MIT Press.')
    add_content(doc, 'Russell, S., Norvig, P. (2020). Artificial Intelligence: A Modern Approach. 4th Edition. Pearson.')
    add_empty_line(doc)
    
    add_section_title(doc, '6.2 Complementaria')
    add_empty_line(doc)
    add_content(doc, 'OpenAI. (2024). GPT-4 Technical Report. Disponible en: https://arxiv.org/abs/2303.08774')
    add_content(doc, 'Anthropic. (2024). Constitutional AI. Disponible en: https://arxiv.org/abs/2212.08073')
    add_content(doc, 'LangChain Documentation. Disponible en: https://docs.langchain.com')
    add_empty_line(doc, 2)
    
    add_section_title(doc, '7. CONOCIMIENTOS PREVIOS EXIGIDOS Y RECOMENDADOS')
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run1 = p.add_run('7.1 Conocimientos Previos Exigidos: ')
    run1.font.name = FONT_NAME
    run1.font.size = FONT_SIZE
    run1.font.bold = True
    run1.font.color.rgb = BLACK
    run2 = p.add_run('Programación Avanzada, Algoritmos y Estructuras de Datos')
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE
    run2.font.bold = False
    run2.font.color.rgb = BLACK
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run1 = p.add_run('7.2 Conocimientos Previos Recomendados: ')
    run1.font.name = FONT_NAME
    run1.font.size = FONT_SIZE
    run1.font.bold = True
    run1.font.color.rgb = BLACK
    run2 = p.add_run('Bases de Datos, Sistemas Operativos')
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE
    run2.font.bold = False
    run2.font.color.rgb = BLACK
    add_empty_line(doc)
    
    add_content(doc, 'No incluye la información de previaturas. Las unidades curriculares previas serán consultadas en el Plan de Estudios de la carrera.')
    add_empty_line(doc, 2)
    
    add_annex_title(doc, 'ANEXO A')
    add_annex_title(doc, 'Para todas las Carreras')
    add_empty_line(doc)
    add_content(doc, 'Esta primera parte del anexo incluye aspectos complementarios que son generales para todas las carreras.')
    add_empty_line(doc)
    
    add_annex_title(doc, 'A1) INSTITUTO')
    add_empty_line(doc)
    add_content(doc, 'Comisión Nacional de Carrera del Tecnólogo en Informática (UTU-UTEC-UDELAR)')
    add_empty_line(doc)
    
    add_annex_title(doc, 'A2) CRONOGRAMA TENTATIVO')
    add_empty_line(doc, 3)
    
    add_annex_title(doc, 'A3) MODALIDAD DEL CURSO Y PROCEDIMIENTO DE EVALUACIÓN')
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run('Aprobación del Curso: ')
    run1.font.name = FONT_NAME
    run1.font.size = FONT_SIZE
    run1.font.bold = True
    run1.font.color.rgb = DARK_GRAY
    run2 = p.add_run('Obtener un promedio ponderado mínimo de 60% en el total de las instancias de evaluación.')
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE
    run2.font.bold = False
    run2.font.color.rgb = DARK_GRAY
    add_empty_line(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run('Presentación Grupal: ')
    run1.font.name = FONT_NAME
    run1.font.size = FONT_SIZE
    run1.font.bold = True
    run1.font.color.rgb = DARK_GRAY
    run2 = p.add_run('Obligatoria para la aprobación del curso.')
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE
    run2.font.bold = False
    run2.font.color.rgb = DARK_GRAY
    add_empty_line(doc)
    
    headers_eval = ['Instancia de Evaluación', 'Peso Relativo', 'Criterio de Aprobación / Descripción']
    rows_eval = [
        ['Parcial 1 (Semana 4)', '40%', 'Presentación Oral Grupal (25 minutos)'],
        ['Evaluación Final (Parcial 2) (Semana 16)', '40%', 'Defensa del Proyecto Integrador'],
        ['Participación y Asistencia', '20%', 'Evaluación de la participación en clase'],
    ]
    create_table_with_header(doc, headers_eval, rows_eval, [Cm(4), Cm(2.5), Cm(10.5)])
    add_empty_line(doc)
    
    add_annex_title(doc, 'A4) CALIDAD DE LIBRE')
    add_empty_line(doc)
    add_content(doc, 'Esta asignatura no adhiere a la resolución del consejo sobre la condición de libre.')
    add_empty_line(doc)
    
    add_annex_title(doc, 'A5) CUPOS DE LA UNIDAD CURRICULAR')
    add_empty_line(doc)
    add_content(doc, 'No tiene cupo')
    add_empty_line(doc, 2)
    
    add_annex_title(doc, 'ANEXO B para la carrera Tecnólogo en Informática')
    add_empty_line(doc)
    
    add_annex_title(doc, 'B1) ÁREA DE FORMACIÓN')
    add_empty_line(doc)
    add_content(doc, 'Inteligencia Artificial y Aprendizaje Automático')
    add_empty_line(doc)
    
    add_annex_title(doc, 'B2) UNIDADES CURRICULARES PREVIAS')
    add_empty_line(doc)
    add_content(doc, 'Programación Avanzada (Examen Aprobado).')
    add_content(doc, 'Algoritmos y Estructuras de Datos (Examen Aprobado).')
    
    output_path = r'C:\utu\utu\ia\01-IA-Ingenieria-Software-Fundamentos\Programa_IA_Ingenieria_Software.docx'
    doc.save(output_path)
    print(f'Documento generado exitosamente: {output_path}')
    return output_path


if __name__ == '__main__':
    create_programa()
