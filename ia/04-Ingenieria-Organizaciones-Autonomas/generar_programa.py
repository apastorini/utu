from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:color'), '000000')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_paragraph_with_style(doc, text, bold=False, font_size=12, font_name='Arial', space_before=0, space_after=0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.alignment = alignment
    return para

def create_program():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_paragraph_with_style(doc, 'Programa de', bold=True, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_style(doc, 'Ingeniería de Organizaciones Autónomas', bold=True, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '1. NOMBRE DE LA UNIDAD CURRICULAR', bold=True, font_size=12)
    doc.add_paragraph()
    add_paragraph_with_style(doc, 'Ingeniería de Organizaciones Autónomas', bold=False, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '2. CRÉDITOS', bold=True, font_size=12)
    doc.add_paragraph()
    add_paragraph_with_style(doc, '8 créditos', bold=False, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '3. OBJETIVOS DE LA UNIDAD CURRICULAR', bold=True, font_size=12)
    doc.add_paragraph()
    
    objetivos = [
        'Construir infraestructuras industriales de agentes con persistencia de memoria a largo plazo',
        'Desplegar sistemas multi-agente escalables mediante contenedores y orquestadores',
        'Optimizar modelos pequeños (SLMs) mediante fine-tuning para tareas empresariales específicas',
        'Implementar Agent-Ops y sistemas de observabilidad',
        'Diseñar arquitecturas de escalamiento horizontal para enjambres de agentes'
    ]
    
    for obj in objetivos:
        para = doc.add_paragraph()
        run = para.add_run(obj)
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        para.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '4. METODOLOGÍA DE ENSEÑANZA', bold=True, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    metodologias = [
        ('Participación de los Estudiantes:', ''),
        ('Demostraciones Prácticas:', 'Despliegue de sistemas multi-agente con Docker y Kubernetes'),
        ('Presentaciones Técnicas (Parcial 1):', 'Arquitectura de orquestación de agentes persistentes'),
        ('Proyecto Final (Evaluación Final):', 'Defensa de infraestructura "Company-in-a-Box" desplegable y escalable')
    ]
    
    for title, desc in metodologias:
        if title == 'Participación de los Estudiantes:':
            para = doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            para.paragraph_format.space_after = Pt(3)
        else:
            para = doc.add_paragraph()
            run = para.add_run(title + ' ')
            run.bold = False
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run2 = para.add_run(desc)
            run2.font.size = Pt(12)
            run2.font.name = 'Arial'
            para.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '5. TEMARIO', bold=True, font_size=12)
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = table.rows[0].cells
    headers[0].text = 'Tema'
    headers[1].text = 'Descripción y subtemas'
    
    for cell in headers:
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    
    temas = [
        ('TEMA I: Workflows Agénticos Industriales', 'Persistencia de estado, gestión de memoria (Redis), patrones de agentes con estado.'),
        ('TEMA II: Integración Legacy', 'Conexión con ERP/CRM, capas de abstracción para agentes, APIs empresariales.'),
        ('TEMA III: Agent-Ops y Observabilidad', 'Tracing de pensamientos, kill-switches semánticos, monitoreo de agentes.'),
        ('TEMA IV: Escalamiento Horizontal', 'Contenerización, enjambres de agentes, fine-tuning de SLMs para tareas específicas.')
    ]
    
    for i, (tema, desc) in enumerate(temas, start=1):
        row = table.rows[i]
        row.cells[0].text = tema
        row.cells[1].text = desc
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '6. BIBLIOGRAFÍA', bold=True, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '6.1 Básica', bold=True, font_size=12)
    doc.add_paragraph()
    
    bibliografia_basica = [
        'Richards, N. (2024). Mathematical Foundations of Reinforcement Learning. Springer.',
        'Sutton, R., Barto, A. (2018). Reinforcement Learning: An Introduction. 2nd Edition. MIT Press.'
    ]
    
    for bib in bibliografia_basica:
        para = doc.add_paragraph()
        para.add_run(bib).font.size = Pt(12)
        para.runs[0].font.name = 'Arial'
        para.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()
    add_paragraph_with_style(doc, '6.2 Complementaria', bold=True, font_size=12)
    doc.add_paragraph()
    
    bibliografia_complementaria = [
        'Docker Documentation. Container Orchestration. Disponible en: https://docs.docker.com',
        'Kubernetes Documentation. Disponible en: https://kubernetes.io/docs',
        'Hugging Face. Fine-tuning Guides. Disponible en: https://huggingface.co/docs/transformers/en/training'
    ]
    
    for bib in bibliografia_complementaria:
        para = doc.add_paragraph()
        para.add_run(bib).font.size = Pt(12)
        para.runs[0].font.name = 'Arial'
        para.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, '7. CONOCIMIENTOS PREVIOS EXIGIDOS Y RECOMENDADOS', bold=True, font_size=12)
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, '7.1 Conocimientos Previos Exigidos:', bold=True, font_size=12)
    para = doc.add_paragraph()
    para.add_run('Programación Avanzada, Sistemas Operativos, IA para la Ingeniería de Software').font.size = Pt(12)
    para.runs[0].font.name = 'Arial'
    para.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    add_paragraph_with_style(doc, '7.2 Conocimientos Previos Recomendados:', bold=True, font_size=12)
    para = doc.add_paragraph()
    para.add_run('Bases de Datos, Redes de Computadoras, Arquitecturas Distribuidas').font.size = Pt(12)
    para.runs[0].font.name = 'Arial'
    para.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'ANEXO A', bold=True, font_size=12)
    add_paragraph_with_style(doc, 'Para todas las Carreras', bold=True, font_size=12)
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Esta primera parte del anexo incluye aspectos complementarios que son generales de la unidad curricular.', bold=False, font_size=12)
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'A1) INSTITUTO', bold=True, font_size=12)
    doc.add_paragraph()
    add_paragraph_with_style(doc, 'Comisión Nacional de Carrera del Tecnólogo en Informática (UTU-UTEC-UDELAR)', bold=False, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'A2) CRONOGRAMA TENTATIVO', bold=True, font_size=12)
    doc.add_paragraph()
    
    cronograma = doc.add_table(rows=17, cols=3)
    cronograma.style = 'Table Grid'
    
    cronograma_headers = cronograma.rows[0].cells
    cronograma_headers[0].text = 'Semana'
    cronograma_headers[1].text = 'Temas y Dominios (4 hs de clase)'
    cronograma_headers[2].text = 'Foco Práctico (2 hs)'
    
    for cell in cronograma_headers:
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    
    cronograma_data = [
        ('Semana 1', 'TEMA I: Introducción a agentes con estado. Conceptos de persistencia.', 'Demostración: Configuración de Redis para memoria de agentes.'),
        ('Semana 2', 'TEMA I: Patrones de workflow agéntico. Máquinas de estado para agentes.', 'Demostración: Implementación de agente con memoria.'),
        ('Semana 3', 'TEMA I: Orquestación multi-agente. Esquemas de comunicación.', 'Demostración: Sistema multi-agente básico.'),
        ('Semana 4', 'TEMA I: Integración con SLMs. Modelos pequeños para producción.', 'Práctica: Fine-tuning de SLM para tarea empresarial.'),
        ('Semana 5', 'TEMA II: Integración con sistemas legacy. ERPs y CRMs.', 'Demostración: Conexión agente-ERP.'),
        ('Semana 6', 'TEMA II: Capas de abstracción. APIs empresariales para agentes.', 'Práctica: Desarrollo de connector de agente.'),
        ('Semana 7', 'TEMA II: Seguridad en integraciones. Autenticación y autorización.', 'Demostración: Agente con acceso seguro a sistemas.'),
        ('Semana 8', 'EVALUACIÓN: Presentaciones Técnicas (Parcial 1)', ''),
        ('Semana 9', 'TEMA III: Agent-Ops. Pipeline de operaciones de agentes.', 'Demostración: Monitoring de agentes en producción.'),
        ('Semana 10', 'TEMA III: Observabilidad. Tracing y logging de pensamientos.', 'Demostración: Implementación de tracing agéntico.'),
        ('Semana 11', 'TEMA III: Kill-switches semánticos. Control y seguridad de IA.', 'Práctica: Implementación de controles de seguridad.'),
        ('Semana 12', 'TEMA IV: Contenerización. Docker para agentes.', 'Demostración: Agente en contenedor Docker.'),
        ('Semana 13', 'TEMA IV: Orquestación con Kubernetes. Escalamiento de agentes.', 'Demostración: Despliegue en Kubernetes.'),
        ('Semana 14', 'TEMA IV: Enjambres de agentes. Patrones de escalamiento horizontal.', 'Tutoría Proyecto Final.'),
        ('Semana 15', 'Revisión General. Proyecto Final - Company-in-a-Box.', 'Tutoría Proyecto Final.'),
        ('Semana 16', 'EVALUACIÓN FINAL: Defensa del Proyecto Company-in-a-Box', '')
    ]
    
    for i, (semana, temas_dom, foco) in enumerate(cronograma_data, start=1):
        row = cronograma.rows[i]
        row.cells[0].text = semana
        row.cells[1].text = temas_dom
        row.cells[2].text = foco
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'A3) MODALIDAD DEL CURSO Y PROCEDIMIENTO DE EVALUACIÓN', bold=True, font_size=12)
    doc.add_paragraph()
    
    tabla_eval = doc.add_table(rows=4, cols=3)
    tabla_eval.style = 'Table Grid'
    
    eval_headers = tabla_eval.rows[0].cells
    eval_headers[0].text = 'Instancia de Evaluación'
    eval_headers[1].text = 'Peso Relativo'
    eval_headers[2].text = 'Criterio de Aprobación / Descripción'
    
    for cell in eval_headers:
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    
    eval_data = [
        ('Parcial 1 (Semana 8)', '40%', 'Presentación Oral Grupal (25 min): Arquitectura de orquestación de agentes persistentes.'),
        ('Evaluación Final (Semana 16)', '40%', 'Defensa del Proyecto Company-in-a-Box: Infraestructura desplegable y escalable.'),
        ('Participación y Asistencia', '20%', 'Evaluación de la participación activa y asistencia a clases.')
    ]
    
    for i, (instancia, peso, criterio) in enumerate(eval_data, start=1):
        row = tabla_eval.rows[i]
        row.cells[0].text = instancia
        row.cells[1].text = peso
        row.cells[2].text = criterio
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'A4) CALIDAD DE LIBRE', bold=True, font_size=12)
    doc.add_paragraph()
    add_paragraph_with_style(doc, 'Esta asignatura no adhiere a la resolución del consejo sobre la condición de libre.', bold=False, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'A5) CUPOS DE LA UNIDAD CURRICULAR', bold=True, font_size=12)
    doc.add_paragraph()
    add_paragraph_with_style(doc, 'No tiene cupo', bold=False, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'ANEXO B para la carrera Tecnólogo en Informática', bold=True, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'B1) ÁREA DE FORMACIÓN', bold=True, font_size=12)
    doc.add_paragraph()
    add_paragraph_with_style(doc, 'Sistemas Inteligentes Avanzados', bold=False, font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph_with_style(doc, 'B2) UNIDADES CURRICULARES PREVIAS', bold=True, font_size=12)
    doc.add_paragraph()
    
    previaturas = [
        'IA para la Ingeniería de Software (Curso Aprobado).',
        'Sistemas Operativos (Curso Aprobado).'
    ]
    
    for prev in previaturas:
        para = doc.add_paragraph()
        para.add_run(prev).font.size = Pt(12)
        para.runs[0].font.name = 'Arial'
        para.paragraph_format.space_after = Pt(3)

    output_path = 'C:/utu/utu/ia/04-Ingenieria-Organizaciones-Autonomas/Programa-Ingenieria-Organizaciones-Autonomas.docx'
    doc.save(output_path)
    print(f'Documento creado exitosamente: {output_path}')

if __name__ == '__main__':
    create_program()
