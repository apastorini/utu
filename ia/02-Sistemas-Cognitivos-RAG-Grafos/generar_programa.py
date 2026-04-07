from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, border_color="000000", border_size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    return p


def set_table_borders(table, color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def create_program():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    title = doc.add_paragraph()
    title_run = title.add_run('Programa de')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(12)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title2 = doc.add_paragraph()
    title2_run = title2.add_run('Sistemas Cognitivos, RAG y Grafos de Conocimiento')
    title2_run.font.name = 'Arial'
    title2_run.font.size = Pt(12)
    title2_run.bold = True
    title2_run.font.color.rgb = RGBColor(0, 0, 0)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_horizontal_line(doc)
    
    def add_section_header(doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True
        return p
    
    def add_normal_text(doc, text, indent=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        if indent:
            p.paragraph_format.left_indent = Inches(0.5)
        return p
    
    def add_bullet_point(doc, text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        return p
    
    add_section_header(doc, '1. NOMBRE DE LA UNIDAD CURRICULAR')
    add_normal_text(doc, 'Sistemas Cognitivos, RAG y Grafos de Conocimiento')
    
    add_section_header(doc, '2. CRÉDITOS')
    add_normal_text(doc, '8 créditos')
    
    add_section_header(doc, '3. OBJETIVOS DE LA UNIDAD CURRICULAR')
    objetivos = [
        'Diseñar arquitecturas de Generación Aumentada por Recuperación (RAG) para reducir alucinaciones',
        'Modelar dominios de conocimiento mediante Ontologías (OWL/RDF) y Bases de Datos de Grafos',
        'Implementar medidas de seguridad avanzada para modelos de lenguaje (OWASP for LLMs)',
        'Integrar percepción computacional con sistemas de recuperación semántica',
        'Construir sistemas multi-agente con orquestación avanzada'
    ]
    for obj in objetivos:
        add_bullet_point(doc, obj)
    
    add_section_header(doc, '4. METODOLOGÍA DE ENSEÑANZA')
    metodologias = [
        ('Demostraciones Prácticas:', 'Implementación de arquitecturas RAG y sistemas de grafos en laboratorio'),
        ('Presentaciones Grupales (Parcial 1):', 'Demostración práctica de arquitectura RAG o de Grafos'),
        ('Proyecto Integrador (Evaluación Final):', 'Defensa de un "Cerebro Cognitivo" integrando recuperación vectorial y reglas ontológicas')
    ]
    for titulo, desc in metodologias:
        p = doc.add_paragraph()
        run1 = p.add_run(titulo + ' ')
        run1.font.name = 'Arial'
        run1.font.size = Pt(12)
        run1.bold = True
        run2 = p.add_run(desc)
        run2.font.name = 'Arial'
        run2.font.size = Pt(12)
    
    add_section_header(doc, '5. TEMARIO')
    
    table1 = doc.add_table(rows=5, cols=3)
    set_table_borders(table1)
    
    headers1 = ['Unidad', 'Temas', 'Descripción']
    header_row = table1.rows[0]
    for i, header in enumerate(headers1):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    temas = [
        ('Unidad I: Percepción Computacional y MLOps', 'Semanas 1-4', 'YOLO, monitoreo de degradación y latencia. Integración de modelos de visión con pipelines de MLops.'),
        ('Unidad II: Arquitectura RAG', 'Semanas 5-8', 'Chunking semántico, bases vectoriales, búsqueda híbrida. Optimización de retrieve-and-generate.'),
        ('Unidad III: Semántica y Grafos', 'Semanas 9-12', 'Modelado OWL, GraphRAG, razonamiento en Neo4j. Construcción de grafos de conocimiento.'),
        ('Unidad IV: Orquestación Multi-Agente', 'Semanas 13-16', 'Prompt Injection, auditoría de decisiones, seguridad LLM. LangChain y AutoGPT.')
    ]
    
    for i, (tema, semanas, desc) in enumerate(temas):
        row = table1.rows[i + 1]
        row.cells[0].text = tema
        row.cells[1].text = semanas
        row.cells[2].text = desc
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    table_hours = doc.add_table(rows=5, cols=3)
    set_table_borders(table_hours)
    
    headers_hours = ['Componente', 'Horas de Clase (Semanales)', 'Horas Totales (16 Semanas)']
    header_row_h = table_hours.rows[0]
    for i, header in enumerate(headers_hours):
        cell = header_row_h.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    horas = [
        ('Teórico (T)', '2 horas', '32 horas'),
        ('Práctico / Laboratorio Asistido (P/L)', '2 horas', '32 horas'),
        ('Dedicación No Presencial (Personal)', '4 horas (estimado)', '64 horas (estimado)'),
        ('Total Estimado', '8 horas', '128 horas')
    ]
    
    for i, (comp, sem, total) in enumerate(horas):
        row = table_hours.rows[i + 1]
        row.cells[0].text = comp
        row.cells[1].text = sem
        row.cells[2].text = total
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if i == 3:
                        run.bold = True
    
    doc.add_paragraph()
    
    table_cronograma = doc.add_table(rows=17, cols=3)
    set_table_borders(table_cronograma)
    
    headers_cron = ['Semana', 'Temas y Dominios (4 hs de clase)', 'Foco Práctico (2 hs)']
    header_row_c = table_cronograma.rows[0]
    for i, header in enumerate(headers_cron):
        cell = header_row_c.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    cronograma = [
        ('Semana 1', 'Unidad I: Introducción a la Percepción Computacional. YOLO y detección de objetos en tiempo real.', 'Demostración: Configuración de entorno YOLO.'),
        ('Semana 2', 'Unidad I: MLOps. Monitoreo de degradación y latencia de modelos.', 'Práctica: Dashboard de monitoreo con Prometheus/Grafana.'),
        ('Semana 3', 'Unidad I: Pipelines de datos. ETL para modelos de visión. Versionamiento de datasets.', 'Práctica: DVC para versionado de datos.'),
        ('Semana 4', 'Unidad I: Fine-tuning y transfer learning. Optimización de modelos para producción.', 'Demostración: Fine-tuning de YOLO.'),
        ('Semana 5', 'Unidad II: Fundamentos RAG. Retrieval-Augmented Generation. Reducción de alucinaciones.', 'Demostración: Implementación básica de RAG con LangChain.'),
        ('Semana 6', 'Unidad II: Chunking semántico. Estrategias de fragmentación de documentos.', 'Práctica: Comparación de estrategias de chunking.'),
        ('Semana 7', 'Unidad II: Bases de datos vectoriales. FAISS, ChromaDB, Pinecone.', 'Demostración: Búsqueda semántica con embeddings.'),
        ('Semana 8', 'Unidad II: Búsqueda híbrida. Combinar búsqueda vectorial y keyword.', 'EVALUACIÓN Parcial 1: Presentaciones Grupales.'),
        ('Semana 9', 'Unidad III: Grafos de conocimiento. RDF, OWL, SPARQL.', 'Demostración: Creación de ontologías con Protégé.'),
        ('Semana 10', 'Unidad III: Neo4j y Cypher. Modelado de grafos de conocimiento.', 'Práctica: Queries Cypher para recuperación semántica.'),
        ('Semana 11', 'Unidad III: GraphRAG. Integración de grafos con RAG tradicional.', 'Demostración: GraphRAG con LangChain y Neo4j.'),
        ('Semana 12', 'Unidad III: Razonamiento ontológico. Reglas y inferencia en grafos.', 'Práctica: Razonamiento con OWL y reglas SWRL.'),
        ('Semana 13', 'Unidad IV: Sistemas Multi-Agente. Orquestación con LangChain Agents.', 'Demostración: Agente RAG con tool calling.'),
        ('Semana 14', 'Unidad IV: Seguridad LLM. OWASP Top 10 for LLMs. Prompt Injection.', 'Práctica: Auditoría de vulnerabilidades LLM.'),
        ('Semana 15', 'Unidad IV: Ética y gobernanza de IA. Sesgos, transparencia, explicabilidad.', 'Tutoría Proyecto Final.'),
        ('Semana 16', 'EVALUACIÓN FINAL: Defensa del "Cerebro Cognitivo" integrando RAG, Grafos y Agentes.', '')
    ]
    
    for i, (semana, temas_dom, foco) in enumerate(cronograma):
        row = table_cronograma.rows[i + 1]
        row.cells[0].text = semana
        row.cells[1].text = temas_dom
        row.cells[2].text = foco
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    add_section_header(doc, '6. BIBLIOGRAFÍA')
    
    p_basica = doc.add_paragraph()
    run = p_basica.add_run('6.1 Básica')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    
    bibliografia_basica = [
        'Gao, Y. et al. (2024). Retrieval-Augmented Generation for Large Language Models: A Survey. arXiv:2312.10997',
        'Hogan, A. et al. (2021). Knowledge Graphs. Morgan & Claypool Publishers.'
    ]
    for bib in bibliografia_basica:
        add_bullet_point(doc, bib)
    
    p_complementaria = doc.add_paragraph()
    run = p_complementaria.add_run('6.2 Complementaria')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    
    bibliografia_complementaria = [
        'LangChain Documentation. RAG Tutorials. Disponible en: https://docs.langchain.com/docs/tutorials/rag',
        'Neo4j Graph Data Science. Disponible en: https://neo4j.com/docs/graph-data-science/current',
        'OWASP. (2024). LLM AI Security and Governance. Disponible en: https://owasp.org/www-project-llmai-security'
    ]
    for bib in bibliografia_complementaria:
        add_bullet_point(doc, bib)
    
    add_section_header(doc, '7. CONOCIMIENTOS PREVIOS EXIGIDOS Y RECOMENDADOS')
    
    p_exigidos = doc.add_paragraph()
    run = p_exigidos.add_run('7.1 Conocimientos Previos Exigidos: ')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    run2 = p_exigidos.add_run('Programación Avanzada, IA para la Ingeniería de Software')
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)
    
    p_recomendados = doc.add_paragraph()
    run = p_recomendados.add_run('7.2 Conocimientos Previos Recomendados: ')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    run2 = p_recomendados.add_run('Bases de Datos, Redes Neuronales Básicas')
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)
    
    add_horizontal_line(doc)
    
    add_section_header(doc, 'ANEXO A')
    p_anexo_a_titulo = doc.add_paragraph()
    run = p_anexo_a_titulo.add_run('Para todas las Carreras')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.italic = True
    
    add_section_header(doc, 'A1) INSTITUTO')
    add_normal_text(doc, 'Comisión Nacional de Carrera del Tecnólogo en Informática (UTU-UTEC-UDELAR)')
    
    add_section_header(doc, 'A2) CRONOGRAMA TENTATIVO')
    table_cron_empty = doc.add_table(rows=1, cols=1)
    set_table_borders(table_cron_empty)
    cell = table_cron_empty.rows[0].cells[0]
    cell.text = ''
    
    add_section_header(doc, 'A3) MODALIDAD DEL CURSO Y PROCEDIMIENTO DE EVALUACIÓN')
    
    table_eval = doc.add_table(rows=5, cols=3)
    set_table_borders(table_eval)
    
    headers_eval = ['Instancia de Evaluación', 'Peso Relativo', 'Criterio de Aprobación / Descripción']
    header_row_eval = table_eval.rows[0]
    for i, header in enumerate(headers_eval):
        cell = header_row_eval.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    evaluaciones = [
        ('Parcial 1 (Semana 8)', '40%', 'Presentación Oral Grupal (25 min): Demostración práctica de arquitectura RAG o de Grafos de Conocimiento.'),
        ('Evaluación Final (Semana 16)', '40%', 'Defensa del "Cerebro Cognitivo": Proyecto integrador que combina recuperación vectorial, reglas ontológicas y orquestación de agentes.'),
        ('Participación y Asistencia', '20%', 'Evaluación de la participación activa y asistencia a clases.'),
        ('Aprobación del Curso', '', 'Obtener un promedio ponderado mínimo de 60% en el total de las instancias de evaluación. La presentación grupal es obligatoria para aprobar el curso.')
    ]
    
    for i, (inst, peso, criterio) in enumerate(evaluaciones):
        row = table_eval.rows[i + 1]
        row.cells[0].text = inst
        row.cells[1].text = peso
        row.cells[2].text = criterio
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    add_section_header(doc, 'A4) CALIDAD DE LIBRE')
    add_normal_text(doc, 'Esta asignatura no adhiere a la resolución del consejo sobre la condición de libre.')
    
    add_section_header(doc, 'A5) CUPOS DE LA UNIDAD CURRICULAR')
    add_normal_text(doc, 'No tiene cupo')
    
    add_horizontal_line(doc)
    
    add_section_header(doc, 'ANEXO B para la carrera Tecnólogo en Informática')
    
    add_section_header(doc, 'B1) ÁREA DE FORMACIÓN')
    add_normal_text(doc, 'Inteligencia Artificial y Aprendizaje Automático')
    
    add_section_header(doc, 'B2) UNIDADES CURRICULARES PREVIAS')
    p_previa = doc.add_paragraph()
    run = p_previa.add_run('IA para la Ingeniería de Software (Curso Aprobado).')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    p_nota = doc.add_paragraph()
    run = p_nota.add_run('No incluye la información de previaturas. Las unidades curriculares previas serán definidas por cada carrera que tome la unidad curricular y serán incluidas en el anexo B.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.italic = True
    
    doc.save(r'C:/utu/utu/ia/02-Sistemas-Cognitivos-RAG-Grafos/Programa_Sistemas_Cognitivos_RAG_Grafos.docx')
    print('Documento generado exitosamente: Programa_Sistemas_Cognitivos_RAG_Grafos.docx')


if __name__ == '__main__':
    create_program()
