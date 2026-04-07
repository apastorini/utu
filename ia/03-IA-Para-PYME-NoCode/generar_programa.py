from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_paragraph_with_format(doc, text, bold=False, font_size=12, font_name='Arial', space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def create_program_document():
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    p = doc.add_paragraph()
    run = p.add_run("Programa de  IA para Líderes y Dueños de PYME (No-Code)")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
    
    add_paragraph_with_format(doc, "1. NOMBRE DE LA UNIDAD CURRICULAR", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "IA para Líderes y Dueños de PYME (No-Code)", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "2. CRÉDITOS ", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "8 créditos", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "3. OBJETIVOS DE LA UNIDAD CURRICULAR ", bold=True, space_after=6)
    doc.add_paragraph()
    
    objetivos = [
        "Identificar oportunidades de automatización agéntica para reducir costos marginales en PYMEs",
        "Orquestar flujos de trabajo de ventas, marketing y administración mediante herramientas visuales (No-Code)",
        "Establecer protocolos de gobernanza y \"Human-in-the-loop\" para la supervisión de agentes",
        "Implementar agentes de front-office para atención 24/7",
        "Construir dashboards de mando y métricas de ROI para monitoreo de agentes"
    ]
    
    for obj in objetivos:
        add_paragraph_with_format(doc, obj, space_after=6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "4. METODOLOGÍA DE ENSEÑANZA", bold=True, space_after=6)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "Participación de los Estudiantes:", space_after=6)
    
    metologias = [
        "Demostraciones Prácticas: Configuración de flujos de automatización con n8n/Make en laboratorio",
        "Seminarios (Parcial 1): Presentación sobre automatización de una unidad de negocio específica",
        "Proyecto Final (Evaluación Final): Defensa de una \"Empresa Autónoma\" con flujo operativo de punta a punta"
    ]
    
    for met in metologias:
        add_paragraph_with_format(doc, met, space_after=6)
    
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "5. TEMARIO", bold=True, space_after=6)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Semanas', 'Unidad Temática', 'Contenido Principal']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
    temario_data = [
        ('Semanas 1-4', 'Economía de la IA', 'Mapeo de flujos de valor, automatización visual (n8n/Make)'),
        ('Semanas 5-8', 'Agentes de Front-Office', 'Vendedores autónomos 24/7, marketing de respuesta directa'),
        ('Semanas 9-12', 'Administración Virtual', 'Lectura de facturas, gestión CRM, conciliación automática'),
        ('Semanas 13-16', 'Dashboards de Mando', 'Control de salud de agentes, métricas de ROI, gobernanza')
    ]
    
    for i, (semanas, unidad, contenido) in enumerate(temario_data):
        row = table.rows[i + 1]
        row.cells[0].text = semanas
        row.cells[1].text = unidad
        row.cells[2].text = contenido
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "6. BIBLIOGRAFÍA", bold=True, space_after=6)
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "6.1 Básica", bold=True, space_after=6)
    doc.add_paragraph()
    
    basica = [
        "Davenport, T., Kirby, J. (2016). Only Humans Need Apply: Winners and Losers in the Age of Smart Machines. Harper Business.",
        "Brynjolfsson, E., McAfee, A. (2014). The Second Machine Age: Work, Progress, and Prosperity. W.W. Norton."
    ]
    
    for bib in basica:
        add_paragraph_with_format(doc, bib, space_after=6)
    
    doc.add_paragraph()
    add_paragraph_with_format(doc, "6.2 Complementaria", bold=True, space_after=6)
    doc.add_paragraph()
    
    complementaria = [
        "n8n Documentation. Workflow Automation. Disponible en: https://docs.n8n.io",
        "Make (formerly Integromat) Documentation. Disponible en: https://www.make.com/en/help",
        "Anthropic. AI Best Practices for Business. Disponible en: https://docs.anthropic.com"
    ]
    
    for bib in complementaria:
        add_paragraph_with_format(doc, bib, space_after=6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "7. CONOCIMIENTOS PREVIOS EXIGIDOS Y RECOMENDADOS", bold=True, space_after=6)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "7.1 Conocimientos Previos Exigidos: Ninguno (curso introductorio para líderes)", space_after=6)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "7.2 Conocimientos Previos Recomendados: Conocimientos básicos de negocio, gestión empresarial", space_after=6)
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "ANEXO A ", bold=True, space_after=6)
    add_paragraph_with_format(doc, "Para todas las Carreras", space_after=6)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "Esta primera parte del anexo incluye aspectos complementarios que son generales de la unidad curricular.", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "A1) INSTITUTO", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "Comisión Nacional de Carrera del Tecnólogo en Informática (UTU-UTEC-UDELAR)", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "A2) CRONOGRAMA TENTATIVO", bold=True, space_after=6)
    doc.add_paragraph()
    
    cronograma_table = doc.add_table(rows=5, cols=3)
    cronograma_table.style = 'Table Grid'
    cronograma_headers = ['Componente', 'Horas de Clase (Semanales)', 'Horas Totales (16 Semanas)']
    for i, header in enumerate(cronograma_headers):
        cell = cronograma_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
    cronograma_data = [
        ('Teórico (T)', '2 horas', '32 horas'),
        ('Práctico / Laboratorio Asistido (P/L)', '2 horas', '32 horas'),
        ('Dedicación No Presencial (Personal)', '4 horas (estimado)', '64 horas (estimado)'),
        ('Total Estimado', '8 horas', '128 horas')
    ]
    
    for i, (comp, semanal, total) in enumerate(cronograma_data):
        row = cronograma_table.rows[i + 1]
        row.cells[0].text = comp
        row.cells[1].text = semanal
        row.cells[2].text = total
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "A3) MODALIDAD DEL CURSO Y PROCEDIMIENTO DE EVALUACIÓN", bold=True, space_after=6)
    doc.add_paragraph()
    doc.add_paragraph()
    
    eval_table = doc.add_table(rows=4, cols=3)
    eval_table.style = 'Table Grid'
    eval_headers = ['Instancia de Evaluación', 'Peso Relativo', 'Criterio de Aprobación / Descripción']
    for i, header in enumerate(eval_headers):
        cell = eval_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
    eval_data = [
        ('Seminario Grupal (Semana 8)', '40%', 'Presentación oral grupal (25 min): Automatización de una unidad de negocio específica'),
        ('Proyecto Final (Semana 16)', '40%', 'Defensa de \"Empresa Autónoma\" con flujo operativo de punta a punta'),
        ('Participación y Asistencia', '20%', 'Evaluación de la participación activa y asistencia')
    ]
    
    for i, (instancia, peso, desc) in enumerate(eval_data):
        row = eval_table.rows[i + 1]
        row.cells[0].text = instancia
        row.cells[1].text = peso
        row.cells[2].text = desc
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "Aprobaci\u00f3n del Curso: Obtener un promedio ponderado m\u00ednimo de 60% en el total de las instancias de evaluaci\u00f3n. La presentaci\u00f3n grupal es obligatoria.", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "A4) CALIDAD DE LIBRE", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "Esta asignatura no adhiere a la resoluci\u00f3n del consejo sobre la condici\u00f3n de libre.", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "A5) CUPOS DE LA UNIDAD CURRICULAR", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "No tiene cupo", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "ANEXO B para la carrera Tecn\u00f3logo en Inform\u00e1tica", bold=True, space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "B1) \u00c1REA DE FORMACI\u00d3N", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "Aplicaciones de Inteligencia Artificial", space_after=12)
    doc.add_paragraph()
    
    add_paragraph_with_format(doc, "B2) UNIDADES CURRICULARES PREVIAS", bold=True, space_after=6)
    doc.add_paragraph()
    add_paragraph_with_format(doc, "Ninguna (puede tomarse en cualquier momento)", space_after=12)
    
    return doc

if __name__ == "__main__":
    doc = create_program_document()
    output_path = "C:/utu/utu/ia/03-IA-Para-PYME-NoCode/Programa IA para Líderes y Dueños de PYME (No-Code).docx"
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")
